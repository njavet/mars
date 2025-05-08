import os
import toml
from pathlib import Path
import re
from docx import Document
from docx.oxml.ns import qn

# project imports
from mars.conf.conf import SYSTEM_PROMPTS, RESULTS_DIR, ALLOWED_HEADINGS


def load_system_messages(filepath=SYSTEM_PROMPTS):
    system_prompts = toml.load(filepath)
    lst = []
    for key, attrs in system_prompts.items():
        s = ''
        for k, v in attrs.items():
            s += '{}: {}\n'.format(k, v)

        dix = {'key': key,
               'text': s}
        lst.append(dix)

    return lst


def format_as_markdown(text: str) -> str:
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.replace('\n', '  \n')


def read_docx(docx_path: Path):
    doc = Document(docx_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text


def get_number_of_runs():
    runs = len([d for d in os.listdir(RESULTS_DIR)
                if os.path.isdir(os.path.join(RESULTS_DIR, d))])
    return runs


def strip_headers_footers(doc: Document) -> None:
    for sect in doc.sections:
        for tag in ('headerReference', 'footerReference'):
            for ref in sect._sectPr.findall(qn(f'w:{tag}')):
                sect._sectPr.remove(ref)
        for part in (
            sect.header, sect.first_page_header, sect.even_page_header,
            sect.footer, sect.first_page_footer, sect.even_page_footer
        ):
            part._element.clear()


def clean_medical_body(doc) -> dict[str, list[str]]:
    strip_headers_footers(doc)

    out: dict[str, list[str]] = {}
    current = None

    for p in doc.paragraphs:
        if p.part is not doc.part:
            continue
        text = p.text.strip()
        if not text:
            continue

        # new section?
        if text in ALLOWED_HEADINGS:
            current = text
            out[current] = []
            continue

        if current:
            out[current].append(text)

    return {k: v for k, v in out.items() if v}
