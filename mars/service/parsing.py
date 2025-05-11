from pathlib import Path
import re
from docx import Document
from docx.oxml.ns import qn
from langchain.text_splitter import RecursiveCharacterTextSplitter

# project imports
from mars.conf import conf


def parse_system_message(sm: str) -> str:
    # remove leading and training ws
    sm = '\n'.join(line.strip() for line in sm.splitlines())
    # replace tabs
    sm = re.sub(r'\t+', ' ', sm)
    # replace multiple ws with one ws
    sm = re.sub(r'[ ]{2,}', ' ', sm)
    # replace 3+ newlines with 2
    sm = re.sub(r'\n{3,}', '\n\n', sm)
    # replace non semantic newlines
    sm = re.sub(r'(?<!\n)\n(?!\n)', ' ', sm)
    # restore bullet point list
    sm = re.sub(r'(?<!\n)\* ', r'\n* ', sm)
    sm = '\n'.join(line.strip() for line in sm.splitlines())
    return sm


def extract_docx_text(docx_dir: Path = conf.DOCX_DIR) -> None:
    for docx_path in docx_dir.glob('*.docx'):
        text = extract_text_from_docx(docx_path)
        out_file = (conf.TEXT_DIR / docx_path.stem).with_suffix('.txt')
        with open(out_file, 'w') as f:
            f.write(text)


def extract_text_from_docx(docx_path: Path) -> str:
    doc = Document(docx_path)
    return '\n'.join([p.text.strip() for p in doc.paragraphs])


def split_docx(dox_path: Path) -> list[str]:
    text = extract_text_from_docx(dox_path)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1024,
        chunk_overlap=0,
        separators=conf.SEPARATORS
    )
    chunks = splitter.split_text(text)
    return chunks


def strip_headers_footers(doc: Document) -> None:
    for sect in doc.sections:
        for tag in ('headerReference', 'footerReference'):
            for ref in sect._sectPr.findall(qn(f'w:{tag}')):
                sect._sectPr.remove(ref)
        for part in (
            sect.header, sect.first_page_header, sect.even_page_header,
            sect.footer, sect.first_page_footer, sect.even_page_footer
        ):
            part._element.clear()


def clean_medical_body(doc) -> dict[str, list[str]]:
    strip_headers_footers(doc)

    out: dict[str, list[str]] = {}
    current = None

    for p in doc.paragraphs:
        if p.part is not doc.part:
            continue
        text = p.text.strip()
        if not text:
            continue

        # new section ?
        if text in conf.ALLOWED_HEADINGS:
            current = text
            out[current] = []
            continue

        if current:
            out[current].append(text)

    return {k: v for k, v in out.items() if v}
