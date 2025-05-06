import io
import json
import os
from fastapi.responses import JSONResponse
from docx import Document
from fastapi import (APIRouter,
                     UploadFile,
                     File,
                     Request,
                     Query,
                     Form)

# project imports
from mars.conf import RESULTS_DIR
from mars.schemas import QueryRequest
from mars.utils.helpers import load_prompts
from mars.utils.helpers import format_as_markdown
from mars.service.rag_context import app_context
from mars.service.service import (get_lms,
                                  run_baseline,
                                  run_baseline_rag)


router = APIRouter()


@router.get('/api/lms')
async def lms(base_url: str = Query(...)):
    return get_lms(base_url)


@router.get('/api/system-messages')
async def get_system_messages() -> JSONResponse:
    return load_prompts()


@router.post("/api/save-scores")
async def save_scores(req: Request):
    data = await req.json()
    file = data['file']
    lm_name = data['lm_name']
    scores = data['scores']

    path = f"frontend/public/results/{file}"
    print(f"saving scores to {path}")
    if not os.path.exists(path):
        return {"error": "File not found"}

    with open(path) as f:
        entries = json.load(f)

    for entry in entries:
        if entry["lm_name"] == lm_name:
            entry["scores"] = scores
            break

    with open(path, "w") as f:
        json.dump(entries, f, indent=2)

    return {"status": "ok"}


@router.post('/api/baseline/base')
async def baseline(payload: QueryRequest) -> JSONResponse:
    res = run_baseline(base_url=payload.server,
                       lm_name=payload.lm_name,
                       system_message=payload.system_message,
                       query=payload.query)
    return JSONResponse({'response': format_as_markdown(res)})


@router.post('/api/baseline/rag')
async def baseline_rag(payload: QueryRequest) -> JSONResponse:
    rag = app_context.rag
    res = run_baseline_rag(base_url=payload.server,
                           lm_name=payload.lm_name,
                           system_message=payload.system_message,
                           query=payload.query,
                           rag=rag)
    return JSONResponse({'response': format_as_markdown(res)})


@router.post('/api/baseline/base-docx')
async def upload_docx(file: UploadFile = File(...),
                      server: str = Form(...),
                      lm_name: str = Form(...),
                      system_message: str = Form(...)) -> JSONResponse:
    contents = await file.read()
    doc = Document(io.BytesIO(contents))
    text = '\n'.join([para.text for para in doc.paragraphs])
    res = run_baseline(base_url=server,
                       lm_name=lm_name,
                       system_message=system_message,
                       query=text)
    return JSONResponse({'response': format_as_markdown(res)})


@router.post('/api/baseline/rag-docx')
async def upload_docx(file: UploadFile = File(...),
                      server: str = Form(...),
                      lm_name: str = Form(...),
                      system_message: str = Form(...)) -> JSONResponse:
    contents = await file.read()
    doc = Document(io.BytesIO(contents))
    text = '\n'.join([para.text for para in doc.paragraphs])
    rag = app_context.rag
    res = run_baseline_rag(base_url=server,
                           lm_name=lm_name,
                           system_message=system_message,
                           query=text,
                           rag=rag)
    return JSONResponse({'response': format_as_markdown(res)})


@router.get('/api/results/file-list')
def list_results():
    files = os.listdir(RESULTS_DIR)
    return files
